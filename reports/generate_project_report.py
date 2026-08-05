from __future__ import annotations

import argparse
import json
import math
from pathlib import Path

import pandas as pd


ROOT = Path(__file__).resolve().parents[1]
REPORT_DIR = ROOT / "reports"
ASSET_DIR = REPORT_DIR / "report_assets"
OUTPUT_DOCX = REPORT_DIR / "Bao_cao_NCKH.docx"


def load_json(path: Path) -> dict:
    with path.open("r", encoding="utf-8") as file:
        return json.load(file)


def metric_row(result: dict) -> dict:
    attack = result["classification_report"]["Attack (1)"]
    return {
        "accuracy": float(result["accuracy"]),
        "precision": float(attack["precision"]),
        "recall": float(attack["recall"]),
        "f1": float(attack["f1-score"]),
        "auc": result.get("auc_roc"),
        "threshold": float(result["decision_threshold"]),
        "cm": result["confusion_matrix"],
    }


def collect_data() -> dict:
    preprocessing = load_json(
        ROOT / "lstm_only" / "artifacts_lstm_only_by_dataset" / "preprocessing_metadata.json"
    )
    proposed_obfu = load_json(
        ROOT
        / "cnn_lstm"
        / "artifacts_cnn_lstm_by_dataset"
        / "by_dataset"
        / "obfu_http"
        / "metadata_and_results.json"
    )
    analysis_summary = load_json(
        ROOT
        / "cnn_lstm"
        / "artifacts_cnn_lstm_by_dataset"
        / "analysis"
        / "obfu_http"
        / "analysis_summary.json"
    )
    comparison_path = (
        ROOT
        / "cnn_lstm"
        / "artifacts_cnn_lstm_by_dataset"
        / "dl_model_comparison_obfu_testbed.csv"
    )
    comparison = pd.read_csv(comparison_path)
    if "Threshold" not in comparison.columns or not comparison["Threshold"].eq(0.5).all():
        raise RuntimeError(
            "Bảng so sánh chưa được re-score đồng nhất tại threshold 0.5. "
            "Hãy chạy cell cuối cnn_lstm/CNN_LSTM.ipynb trước khi sinh báo cáo."
        )
    sequential_parallel_path = (
        ROOT
        / "cnn_lstm"
        / "artifacts_cnn_lstm_by_dataset"
        / "sequential_vs_parallel_by_dataset.csv"
    )
    sequential_parallel = pd.read_csv(sequential_parallel_path)
    if not sequential_parallel["Evaluation Threshold"].eq(0.5).all():
        raise RuntimeError(
            "Bảng tuần tự/song song chưa được đánh giá đồng nhất tại threshold 0.5."
        )

    model_paths = {
        "CNN": {
            "kaggle": ROOT / "cnn_only/artifacts_cnn_only_by_dataset/by_dataset/kaggle/metadata_and_results.json",
            "csic": ROOT / "cnn_only/artifacts_cnn_only_by_dataset/by_dataset/csic/metadata_and_results.json",
            "obfu_http": ROOT / "cnn_only/artifacts_cnn_only_by_dataset/by_dataset/obfu_http/metadata_and_results.json",
        },
        "LSTM": {
            "kaggle": ROOT / "lstm_only/artifacts_lstm_only_by_dataset/by_dataset/kaggle/metadata_and_results.json",
            "csic": ROOT / "lstm_only/artifacts_lstm_only_by_dataset/by_dataset/csic/metadata_and_results.json",
            "obfu_http": ROOT / "lstm_only/artifacts_lstm_only_by_dataset/by_dataset/obfu_http/metadata_and_results.json",
        },
        "CNN–LSTM": {
            "kaggle": ROOT / "cnn_lstm/artifacts_cnn_lstm_by_dataset/by_dataset/kaggle/metadata_and_results.json",
            "csic": ROOT / "cnn_lstm/artifacts_cnn_lstm_by_dataset/by_dataset/csic/metadata_and_results.json",
            "obfu_http": ROOT / "cnn_lstm/artifacts_cnn_lstm_by_dataset/by_dataset/obfu_http/metadata_and_results.json",
        },
    }
    models = {
        model: {dataset: load_json(path) for dataset, path in paths.items()}
        for model, paths in model_paths.items()
    }

    obfu_df = pd.read_csv(ROOT / "obfu_http_dataset_v2.csv", low_memory=False)
    obfu_stats = {
        "rows": int(len(obfu_df)),
        "normal": int((obfu_df["classification"] == "normal").sum()),
        "attack": int((obfu_df["classification"] == "anomalous").sum()),
        "sqli": int((obfu_df["attack_category"] == "sqli").sum()),
        "xss": int((obfu_df["attack_category"] == "xss").sum()),
        "seed_count": int(
            obfu_df.loc[obfu_df["classification"] == "anomalous", "seed_id"].nunique()
        ),
        "combo_count": int(obfu_df["assigned_combo"].nunique()),
        "technique_string_count": int(obfu_df["obfuscation_techniques"].nunique()),
        "context": obfu_df["context_location"].fillna("normal/không áp dụng").value_counts().to_dict(),
        "difficulty": obfu_df["difficulty_level"].fillna("normal/không áp dụng").value_counts().to_dict(),
        "obfuscation_type": obfu_df["obfuscation_type"].fillna("normal/không áp dụng").value_counts().to_dict(),
    }
    return {
        "preprocessing": preprocessing,
        "proposed_obfu": proposed_obfu,
        "analysis_summary": analysis_summary,
        "comparison": comparison,
        "sequential_parallel": sequential_parallel,
        "models": models,
        "obfu_stats": obfu_stats,
    }


def generate_assets(data: dict) -> None:
    import matplotlib

    # Use a non-interactive backend so report generation also works on Kaggle,
    # CI, and Windows installations without a complete Tcl/Tk runtime.
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    from matplotlib.patches import FancyArrowPatch, FancyBboxPatch

    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    plt.rcParams.update({"font.family": "DejaVu Sans", "font.size": 10})

    # Dataset distribution.
    datasets = data["preprocessing"]["datasets"]
    names = ["Kaggle\nSQLi–XSS", "HTTP\nCSIC 2010", "OBFU_HTTP\nv2"]
    keys = ["kaggle", "csic", "obfu_http"]
    normal = [datasets[key]["clean"]["label_counts"]["0"] for key in keys]
    attack = [datasets[key]["clean"]["label_counts"]["1"] for key in keys]
    fig, ax = plt.subplots(figsize=(8.4, 4.6))
    x = range(len(keys))
    ax.bar(x, normal, label="Normal (0)", color="#3B82F6")
    ax.bar(x, attack, bottom=normal, label="Attack (1)", color="#EF4444")
    for index, (n0, n1) in enumerate(zip(normal, attack)):
        ax.text(index, n0 / 2, f"{n0:,}", ha="center", va="center", color="white", weight="bold")
        ax.text(index, n0 + n1 / 2, f"{n1:,}", ha="center", va="center", color="white", weight="bold")
    ax.set_xticks(list(x), names)
    ax.set_ylabel("Số mẫu sau làm sạch")
    ax.set_title("Phân bố nhãn của ba nguồn dữ liệu")
    ax.legend(frameon=False, ncol=2, loc="upper center")
    ax.grid(axis="y", alpha=0.2)
    fig.tight_layout()
    fig.savefig(ASSET_DIR / "dataset_distribution.png", dpi=220, bbox_inches="tight")
    plt.close(fig)

    # Architecture diagram.
    fig, ax = plt.subplots(figsize=(13, 3.2))
    ax.set_xlim(0, 13)
    ax.set_ylim(0, 3)
    ax.axis("off")
    stages = [
        (0.2, "Input\n768/1024", "#E0F2FE"),
        (1.55, "Embedding\n64", "#DBEAFE"),
        (2.9, "Conv1D\n128, k=3", "#DCFCE7"),
        (4.25, "MaxPool\n4", "#DCFCE7"),
        (5.6, "Conv1D\n128, k=5", "#DCFCE7"),
        (6.95, "MaxPool\n4", "#DCFCE7"),
        (8.3, "LSTM\n128", "#F3E8FF"),
        (9.65, "Global\nMax Pool", "#F3E8FF"),
        (11.0, "Dense 64\nDropout 0.3", "#FEF3C7"),
        (12.25, "Sigmoid\nP(Attack)", "#FEE2E2"),
    ]
    for x0, label, color in stages:
        width = 1.05 if x0 < 12 else 0.7
        box = FancyBboxPatch(
            (x0, 1.05), width, 0.9,
            boxstyle="round,pad=0.03,rounding_size=0.08",
            facecolor=color, edgecolor="#334155", linewidth=1.2,
        )
        ax.add_patch(box)
        ax.text(x0 + width / 2, 1.5, label, ha="center", va="center", fontsize=9, weight="bold")
    for (x0, _, _), (x1, _, _) in zip(stages, stages[1:]):
        w0 = 1.05 if x0 < 12 else 0.7
        ax.add_patch(FancyArrowPatch((x0 + w0, 1.5), (x1, 1.5), arrowstyle="-|>", mutation_scale=12, color="#475569"))
    ax.text(6.5, 2.55, "Kiến trúc CNN–LSTM tuần tự được triển khai", ha="center", fontsize=14, weight="bold")
    ax.text(5.55, 0.48, "CNN: học motif cục bộ", ha="center", color="#166534", weight="bold")
    ax.text(8.95, 0.48, "LSTM: học ngữ cảnh theo thứ tự", ha="center", color="#6B21A8", weight="bold")
    fig.tight_layout()
    fig.savefig(ASSET_DIR / "cnn_lstm_architecture.png", dpi=220, bbox_inches="tight")
    plt.close(fig)

    # Web application flow.
    fig, ax = plt.subplots(figsize=(11.5, 3.5))
    ax.set_xlim(0, 11.5)
    ax.set_ylim(0, 3.5)
    ax.axis("off")
    flow = [
        (0.3, "Trình duyệt\nHTML/CSS/JS", "#E0F2FE"),
        (2.35, "POST\n/api/predict", "#DBEAFE"),
        (4.4, "Chuẩn hóa\nHTTP envelope", "#DCFCE7"),
        (6.45, "Tokenizer +\npad/truncate", "#FEF3C7"),
        (8.5, "CNN–LSTM\ncheckpoint", "#F3E8FF"),
        (10.4, "JSON\nresult", "#FEE2E2"),
    ]
    for x0, label, color in flow:
        box = FancyBboxPatch((x0, 1.2), 1.45, 1.0, boxstyle="round,pad=0.05", facecolor=color, edgecolor="#334155")
        ax.add_patch(box)
        ax.text(x0 + 0.725, 1.7, label, ha="center", va="center", weight="bold")
    for (x0, _, _), (x1, _, _) in zip(flow, flow[1:]):
        ax.add_patch(FancyArrowPatch((x0 + 1.45, 1.7), (x1, 1.7), arrowstyle="-|>", mutation_scale=12, color="#475569"))
    ax.text(5.75, 3.0, "Luồng suy luận của WebApp Flask", ha="center", fontsize=14, weight="bold")
    fig.tight_layout()
    fig.savefig(ASSET_DIR / "webapp_flow.png", dpi=220, bbox_inches="tight")
    plt.close(fig)

    # Research timeline.
    fig, ax = plt.subplots(figsize=(11.5, 4.2))
    ax.axis("off")
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 4)
    milestones = [
        (0.5, "20/06", "Khởi tạo\nCNN–LSTM + WebApp"),
        (2.2, "21–24/06", "Chọn siêu tham số,\nLSTM-only, báo cáo"),
        (4.0, "01–18/07", "CNN-only, song song,\nretrain baseline"),
        (5.9, "21–24/07", "Chuẩn hóa preprocessing,\nsplit theo nguồn"),
        (7.7, "02/08", "Dataset OBFU v2,\nquality gate, phân tích"),
        (9.3, "05/08", "Đồng bộ threshold,\nWebApp và báo cáo"),
    ]
    ax.plot([0.5, 9.3], [2, 2], color="#64748B", linewidth=2)
    for index, (x0, date, label) in enumerate(milestones):
        ax.scatter([x0], [2], s=120, color="#2563EB", zorder=3)
        above = index % 2 == 0
        y = 2.65 if above else 1.15
        ax.text(x0, y + (0.35 if above else -0.35), date, ha="center", weight="bold", color="#1E3A8A")
        ax.text(x0, y, label, ha="center", va="center", fontsize=9)
    ax.set_title("Tiến trình nghiên cứu và triển khai (theo lịch sử Git)", fontsize=14, weight="bold")
    fig.tight_layout()
    fig.savefig(ASSET_DIR / "research_timeline.png", dpi=220, bbox_inches="tight")
    plt.close(fig)

    # OBFU held-out results.
    rows = []
    for name, result in data["proposed_obfu"]["evaluation"].items():
        label = name.split(":")[-1] if ":" in name else "test"
        metrics = metric_row(result)
        rows.append((label, metrics["precision"] * 100, metrics["recall"] * 100, metrics["f1"] * 100))
    order = {"test": 0, "test_unseen_technique": 1, "test_unseen_seed": 2, "test_unseen_both": 3}
    rows.sort(key=lambda row: order[row[0]])
    fig, ax = plt.subplots(figsize=(9.5, 4.8))
    x = list(range(len(rows)))
    width = 0.24
    for offset, metric_index, label, color in [(-width, 1, "Precision", "#2563EB"), (0, 2, "Recall", "#DC2626"), (width, 3, "F1", "#16A34A")]:
        values = [row[metric_index] for row in rows]
        bars = ax.bar([value + offset for value in x], values, width, label=label, color=color)
        ax.bar_label(bars, fmt="%.2f", padding=2, fontsize=8)
    ax.set_xticks(x, [row[0].replace("test_unseen_", "unseen\n") for row in rows])
    ax.set_ylim(98.8, 100.08)
    ax.set_ylabel("Phần trăm")
    ax.set_title("CNN–LSTM trên các split OBFU_HTTP v2 (threshold 0,5)")
    ax.legend(frameon=False, ncol=3)
    ax.grid(axis="y", alpha=0.2)
    fig.tight_layout()
    fig.savefig(ASSET_DIR / "obfu_generalization.png", dpi=220, bbox_inches="tight")
    plt.close(fig)

    # Fair comparison at threshold 0.5.
    comparison = data["comparison"]
    fig, axes = plt.subplots(1, 3, figsize=(13, 4.4), sharey=True)
    for axis, dataset in zip(axes, ["HTTP CSIC2010", "SQLi - XSS Payload", "Testbed Dataset"]):
        frame = comparison[comparison["Dataset"] == dataset]
        x = list(range(len(frame)))
        width = 0.25
        for offset, column, label, color in [(-width, "Precision", "Precision", "#2563EB"), (0, "Recall", "Recall", "#DC2626"), (width, "F1-Score", "F1", "#16A34A")]:
            bars = axis.bar([value + offset for value in x], frame[column], width, label=label, color=color)
            axis.bar_label(bars, fmt="%.2f", padding=2, fontsize=7, rotation=90)
        axis.set_xticks(x, frame["DL Model"].str.replace("PROPOSED MODEL", "CNN–LSTM"))
        axis.set_title(dataset, pad=26)
        axis.grid(axis="y", alpha=0.2)
    axes[0].set_ylabel("Phần trăm")
    comparison_min = float(
        comparison[["Precision", "Recall", "F1-Score"]].min().min()
    )
    axes[0].set_ylim(max(0, comparison_min - 0.55), 100.55)
    handles, labels = axes[0].get_legend_handles_labels()
    fig.legend(handles, labels, loc="lower center", ncol=3, frameon=False)
    fig.suptitle("So sánh ba kiến trúc tại cùng threshold 0,5", fontsize=14, weight="bold")
    fig.tight_layout(rect=(0, 0.08, 1, 0.93))
    fig.savefig(ASSET_DIR / "fair_model_comparison.png", dpi=220, bbox_inches="tight")
    plt.close(fig)

    # Sequential CNN-LSTM versus parallel CNN || LSTM.
    topology = data["sequential_parallel"]
    fig, axes = plt.subplots(1, 2, figsize=(10.2, 4.6), sharey=True)
    for axis, dataset in zip(axes, ["KAGGLE", "CSIC"]):
        frame = topology[topology["Dataset"] == dataset]
        labels = frame["Architecture"].str.replace("CNN–LSTM tuần tự", "Tuần tự")
        labels = labels.str.replace("CNN ∥ LSTM song song", "Song song")
        x = list(range(len(frame)))
        width = 0.25
        for offset, column, label, color in [
            (-width, "Accuracy", "Accuracy", "#2563EB"),
            (0, "Attack Recall", "Recall", "#DC2626"),
            (width, "Attack F1", "F1", "#16A34A"),
        ]:
            bars = axis.bar(
                [value + offset for value in x], frame[column], width,
                label=label, color=color,
            )
            axis.bar_label(bars, fmt="%.3f", padding=2, fontsize=8, rotation=90)
        axis.set_xticks(x, labels)
        axis.set_title(dataset)
        axis.grid(axis="y", alpha=0.2)
    min_value = float(
        topology[["Accuracy", "Attack Recall", "Attack F1"]].min().min()
    )
    axes[0].set_ylim(max(0, min_value - 0.35), 100.15)
    axes[0].set_ylabel("Phần trăm")
    handles, labels = axes[0].get_legend_handles_labels()
    fig.legend(handles, labels, loc="lower center", ncol=3, frameon=False)
    fig.suptitle(
        "CNN–LSTM tuần tự và CNN ∥ LSTM song song tại threshold 0,5",
        fontsize=14,
        weight="bold",
    )
    fig.tight_layout(rect=(0, 0.09, 1, 0.93))
    fig.savefig(ASSET_DIR / "sequential_vs_parallel.png", dpi=220, bbox_inches="tight")
    plt.close(fig)


def generate_docx(data: dict) -> None:
    from docx import Document
    from docx.enum.section import WD_ORIENT, WD_SECTION
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.opc.constants import RELATIONSHIP_TYPE
    from docx.shared import Cm, Inches, Pt, RGBColor

    def set_cell_shading(cell, fill: str) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = tc_pr.find(qn("w:shd"))
        if shd is None:
            shd = OxmlElement("w:shd")
            tc_pr.append(shd)
        shd.set(qn("w:fill"), fill)

    def set_cell_text(cell, text, bold=False, color=None, size=9) -> None:
        cell.text = ""
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(str(text))
        run.bold = bold
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = RGBColor.from_string(color)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    def add_table(doc, headers, rows, widths=None, proposed_col=None):
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        for index, header in enumerate(headers):
            set_cell_text(table.rows[0].cells[index], header, bold=True, color="FFFFFF", size=9)
            set_cell_shading(table.rows[0].cells[index], "1F4E78")
        for row in rows:
            cells = table.add_row().cells
            is_proposed = proposed_col is not None and "CNN" in str(row[proposed_col]) and "LSTM" in str(row[proposed_col])
            for index, value in enumerate(row):
                set_cell_text(cells[index], value, bold=is_proposed, size=9)
                if is_proposed:
                    set_cell_shading(cells[index], "EAF2F8")
        doc.add_paragraph()
        return table

    def add_caption(doc, text):
        paragraph = doc.add_paragraph(style="Caption VN")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run(text)
        return paragraph

    def add_figure(doc, filename, caption, width=6.7):
        path = ASSET_DIR / filename
        if not path.exists():
            raise FileNotFoundError(f"Thiếu hình báo cáo: {path}")
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run().add_picture(str(path), width=Inches(width))
        add_caption(doc, caption)

    def add_code(doc, text):
        paragraph = doc.add_paragraph(style="Code VN")
        paragraph.paragraph_format.keep_together = True
        paragraph.add_run(text)
        return paragraph

    def add_bullets(doc, items, level=0):
        for item in items:
            paragraph = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
            paragraph.add_run(item)

    def add_numbered(doc, items):
        for item in items:
            paragraph = doc.add_paragraph(style="List Number")
            paragraph.add_run(item)

    def add_hyperlink(paragraph, url, label=None):
        """Add a real clickable external hyperlink to a python-docx paragraph."""
        relationship_id = paragraph.part.relate_to(
            url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True
        )
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), relationship_id)
        run = OxmlElement("w:r")
        run_properties = OxmlElement("w:rPr")
        color = OxmlElement("w:color")
        color.set(qn("w:val"), "0563C1")
        underline = OxmlElement("w:u")
        underline.set(qn("w:val"), "single")
        run_properties.extend([color, underline])
        text_element = OxmlElement("w:t")
        text_element.text = label or url
        run.extend([run_properties, text_element])
        hyperlink.append(run)
        paragraph._p.append(hyperlink)
        return hyperlink

    def add_toc(paragraph):
        run = paragraph.add_run()
        fld_char = OxmlElement("w:fldChar")
        fld_char.set(qn("w:fldCharType"), "begin")
        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'
        fld_sep = OxmlElement("w:fldChar")
        fld_sep.set(qn("w:fldCharType"), "separate")
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run._r.extend([fld_char, instr_text, fld_sep, fld_end])

    def add_page_number(paragraph):
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        instruction = OxmlElement("w:instrText")
        instruction.set(qn("xml:space"), "preserve")
        instruction.text = " PAGE "
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        run._r.extend([begin, instruction, end])

    def pct(value, digits=2):
        return f"{float(value) * 100:.{digits}f}%".replace(".", ",")

    def num(value):
        return f"{int(value):,}".replace(",", ".")

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.0)
    section.header_distance = Cm(1.0)
    section.footer_distance = Cm(1.0)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.first_line_indent = Cm(1.0)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.space_after = Pt(6)
    for style_name, size, color in [("Title", 24, "1F4E78"), ("Heading 1", 16, "1F4E78"), ("Heading 2", 14, "2F5597"), ("Heading 3", 12, "44546A")]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(6)
    caption_style = styles.add_style("Caption VN", WD_STYLE_TYPE.PARAGRAPH)
    caption_style.font.name = "Times New Roman"
    caption_style.font.size = Pt(10)
    caption_style.font.italic = True
    caption_style.paragraph_format.space_after = Pt(8)
    code_style = styles.add_style("Code VN", WD_STYLE_TYPE.PARAGRAPH)
    code_style.font.name = "Consolas"
    code_style.font.size = Pt(9)
    code_style.paragraph_format.left_indent = Cm(0.7)
    code_style.paragraph_format.right_indent = Cm(0.4)
    code_style.paragraph_format.space_before = Pt(4)
    code_style.paragraph_format.space_after = Pt(6)
    code_style.paragraph_format.line_spacing = 1.0

    for section_item in doc.sections:
        header = section_item.header.paragraphs[0]
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = header.add_run("BÁO CÁO DỰ ÁN – PHÁT HIỆN XSS/SQLi BỊ OBFUSCATION")
        run.font.name = "Times New Roman"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(100, 116, 139)
        add_page_number(section_item.footer.paragraphs[0])

    # Cover.
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("[TÊN TRƯỜNG / ĐƠN VỊ]\n").bold = True
    p.add_run("[KHOA / BỘ MÔN]").bold = True
    doc.add_paragraph("\n")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("BÁO CÁO TỔNG KẾT DỰ ÁN")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(31, 78, 120)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ỨNG DỤNG HỌC MÁY ĐỂ PHÁT HIỆN\nPAYLOAD XSS VÀ SQL INJECTION BỊ OBFUSCATION")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(20)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Mô hình character-level CNN–LSTM và WebApp suy luận")
    run.italic = True
    run.font.size = Pt(14)
    doc.add_paragraph("\n\n")
    cover_rows = [
        ("Nhóm thực hiện", "[ĐIỀN TÊN NHÓM / THÀNH VIÊN]"),
        ("Giảng viên hướng dẫn", "[ĐIỀN TÊN GIẢNG VIÊN]"),
        ("Mã đề tài", "[ĐIỀN NẾU CÓ]"),
        ("Thời gian thực hiện", "06/2026 – 08/2026"),
    ]
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for left, right in cover_rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], left, bold=True, size=11)
        set_cell_text(cells[1], right, size=11)
    doc.add_paragraph("\n\n")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Tháng 8 năm 2026").bold = True
    doc.add_page_break()

    # Abstract.
    fixed_comparison = data["comparison"]
    hybrid_rows = fixed_comparison[
        fixed_comparison["DL Model"] == "PROPOSED MODEL"
    ].set_index("Dataset")
    hybrid_kaggle = hybrid_rows.loc["SQLi - XSS Payload"]
    hybrid_csic = hybrid_rows.loc["HTTP CSIC2010"]
    doc.add_heading("TÓM TẮT", level=1)
    doc.add_paragraph(
        "Dự án nghiên cứu bài toán phát hiện request/payload web chứa SQL Injection (SQLi) hoặc Cross-Site Scripting (XSS), đặc biệt khi chuỗi tấn công đã được che giấu bằng mã hóa, thay đổi biểu diễn hoặc kết hợp nhiều kỹ thuật obfuscation. Hệ thống sử dụng biểu diễn character-level, giữ nguyên case và dấu vết encoding, sau đó phân loại nhị phân Normal/Attack bằng mạng CNN–LSTM tuần tự. CNN học các motif cục bộ như ký tự đặc biệt, comment, thẻ HTML và chuỗi mã hóa; LSTM học thứ tự và ngữ cảnh của các motif trên toàn request."
    )
    doc.add_paragraph(
        "Quá trình nghiên cứu không dừng ở việc huấn luyện một mô hình. Nhóm xây dựng ba baseline CNN-only, LSTM-only và CNN–LSTM; phát hiện vấn đề shortcut và rò rỉ seed trong dataset obfuscation thế hệ đầu; thiết kế lại OBFU_HTTP v2 với 129.998 request cân bằng, nhiều vị trí chèn payload và ba tập kiểm tra giữ lại seed/kỹ thuật chưa thấy. Pipeline chung chuẩn hóa mọi đầu vào thành HTTP envelope, tokenizer chỉ fit trên train split và các mô hình được re-score ở threshold 0,5 để so sánh công bằng."
    )
    doc.add_paragraph(
        f"Sau lần huấn luyện đồng bộ mới nhất, CNN–LSTM tuần tự đạt F1 "
        f"{hybrid_kaggle['F1-Score']:.3f}% trên Kaggle SQLi–XSS và "
        f"{hybrid_csic['F1-Score']:.3f}% trên HTTP CSIC 2010 tại threshold 0,5. "
        "Trên OBFU_HTTP test, mô hình đạt accuracy 99,963%, Attack precision "
        "99,956%, recall 99,971% và F1 99,963%; trên test_unseen_both, F1 còn "
        "99,915%. CNN–LSTM được chọn làm mô hình triển khai vì tạo một chuỗi xử lý "
        "phân cấp: CNN trích motif và nén chuỗi trước khi LSTM mô hình hóa thứ tự. "
        "Kết quả đối chứng song song cho thấy lựa chọn topology vẫn phụ thuộc dataset, "
        "nên báo cáo không tuyên bố tuần tự luôn vượt trội."
    )
    p = doc.add_paragraph()
    p.add_run("Từ khóa: ").bold = True
    p.add_run("SQL Injection, XSS, obfuscation, character-level, CNN, LSTM, deep learning, Flask.")
    doc.add_page_break()

    doc.add_heading("MỤC LỤC", level=1)
    toc_paragraph = doc.add_paragraph()
    add_toc(toc_paragraph)
    doc.add_paragraph("Lưu ý: mở file trong Microsoft Word, nhấn Ctrl+A rồi F9 để cập nhật mục lục và số trang.")
    doc.add_page_break()

    doc.add_heading("DANH MỤC TỪ VIẾT TẮT", level=1)
    add_table(doc, ["Từ viết tắt", "Diễn giải"], [
        ("SQLi", "SQL Injection"), ("XSS", "Cross-Site Scripting"),
        ("CNN", "Convolutional Neural Network"), ("LSTM", "Long Short-Term Memory"),
        ("DL", "Deep Learning"), ("HTTP", "Hypertext Transfer Protocol"),
        ("API", "Application Programming Interface"), ("AUC", "Area Under the Curve"),
        ("FPR/FNR", "False Positive Rate / False Negative Rate"), ("OOV", "Out-of-vocabulary"),
    ])

    doc.add_heading("1. GIỚI THIỆU", level=1)
    doc.add_heading("1.1. Bối cảnh và lý do chọn đề tài", level=2)
    doc.add_paragraph(
        "SQL Injection và XSS là hai nhóm tấn công injection phổ biến trên ứng dụng web. SQLi đưa lệnh SQL ngoài ý muốn vào dữ liệu đầu vào nhằm làm thay đổi truy vấn; XSS đưa nội dung có thể thực thi phía trình duyệt vào luồng dữ liệu của ứng dụng. OWASP mô tả cả hai là các rủi ro có thể ảnh hưởng đến bí mật, toàn vẹn dữ liệu, xác thực, phiên người dùng và nội dung hiển thị [1], [2]."
    )
    doc.add_paragraph(
        "Khó khăn thực tế là payload hiếm khi xuất hiện ở dạng rõ. Kẻ tấn công có thể dùng URL encoding, HTML entities, Unicode escape, hex, thay đổi hoa/thường, comment, chèn ký tự phân tách, double encoding hoặc tổ hợp nhiều phép biến đổi. OWASP Web Security Testing Guide ghi nhận encoded injection có thể làm chuỗi độc hại vượt qua bộ lọc đầu vào yếu [3]. Vì vậy detector dựa thuần vào danh sách từ khóa hoặc regex dễ bỏ sót biến thể chưa được mô tả trước."
    )
    doc.add_heading("1.2. Phát biểu bài toán", level=2)
    doc.add_paragraph(
        "Cho đầu vào x là payload, URL hoặc HTTP request. Sau khi chuyển về biểu diễn thống nhất, mô hình ước lượng xác suất ŷ = P(y=1|x), trong đó y=1 là Attack và y=0 là Normal. Với threshold t=0,5, request được gán Attack khi ŷ ≥ t. Bài toán hiện tại là phân loại nhị phân; SQLi và XSS cùng thuộc lớp Attack vì mục tiêu triển khai là cảnh báo request nguy hiểm trước khi phân loại chuyên sâu."
    )
    doc.add_heading("1.3. Mục tiêu", level=2)
    add_numbered(doc, [
        "Nghiên cứu đặc trưng của SQLi, XSS và các kỹ thuật che giấu payload.",
        "Xây dựng pipeline dữ liệu giữ nguyên dấu vết obfuscation và hạn chế leakage.",
        "Xây dựng các baseline CNN-only, LSTM-only và mô hình CNN–LSTM.",
        "Đánh giá in-domain, cross-source và trên các split giữ lại seed/kỹ thuật chưa thấy.",
        "Triển khai checkpoint thành WebApp Flask có giao diện và API suy luận.",
        "Ghi nhận trung thực giới hạn của dataset, thiết kế thí nghiệm và khả năng khái quát.",
    ])
    doc.add_heading("1.4. Phạm vi và đóng góp", level=2)
    add_bullets(doc, [
        "Đầu vào: chuỗi ký tự đại diện cho payload, URL, raw HTTP request hoặc HTTP envelope.",
        "Đầu ra: xác suất Attack, nhãn Normal/Attack và thông tin preprocessing.",
        "Ba nguồn dữ liệu được train thành các model độc lập; không ghép ba nguồn thành một train set duy nhất.",
        "Đóng góp chính: OBFU_HTTP v2, pipeline character-level chung, kiến trúc lai tuần tự, baseline ablation và WebApp end-to-end.",
        "Hệ thống là công cụ hỗ trợ phát hiện, không thay thế parameterized query, output encoding, CSP, validation và các biện pháp secure coding.",
    ])

    doc.add_heading("2. CƠ SỞ LÝ THUYẾT", level=1)
    doc.add_heading("2.1. SQL Injection", level=2)
    doc.add_paragraph(
        "SQLi xảy ra khi dữ liệu không tin cậy được ghép vào câu truy vấn SQL động, khiến dữ liệu và lệnh điều khiển không còn được tách biệt. Biến thể trong dataset gồm auth bypass, UNION-based, error-based, Boolean blind, time blind, stacked query, out-of-band, second-order và no-separator. Obfuscation có thể làm thay đổi bề mặt chuỗi nhưng vẫn bảo toàn ý nghĩa đối với DBMS."
    )
    doc.add_heading("2.2. Cross-Site Scripting", level=2)
    doc.add_paragraph(
        "XSS đưa HTML/JavaScript hoặc URI có khả năng thực thi vào ngữ cảnh trình duyệt. Dataset bao phủ script tag, event handler, SVG/HTML5 context, URI scheme, CSS context, tag breakout, mutation XSS, blind và stored XSS. Khả năng có nhiều context khiến một luật lọc duy nhất khó bao phủ đầy đủ."
    )
    doc.add_heading("2.3. Obfuscation và tác động đến detector", level=2)
    add_bullets(doc, [
        "Thay đổi mã hóa: percent/URL encoding, HTML entity, Unicode/hex/char encoding.",
        "Thay đổi từ vựng: mixed case, comment, whitespace, separator hoặc nối chuỗi.",
        "Thay đổi ngữ cảnh vận chuyển: query, form, JSON body, cookie, header, path, multipart.",
        "Kết hợp nhiều kỹ thuật tạo chuỗi chưa xuất hiện nguyên văn trong train set.",
        "Double encoding hoặc normalization không thống nhất có thể tạo khoảng trống giữa detector và thành phần thực thi.",
    ])
    doc.add_paragraph(
        "Nguy cơ này không chỉ mang tính giả định. WAF-A-MoLE tạo các phép biến đổi "
        "bảo toàn ngữ nghĩa để né bộ phát hiện SQLi, còn AdvSQLi tiếp tục chứng minh "
        "các chuỗi SQLi đối kháng có thể vượt qua cả detector học máy và WAF thương mại "
        "trong thiết lập hộp đen [14], [15]. Hai công trình là cơ sở để dự án xem "
        "obfuscation và unseen transformation như yêu cầu đánh giá chính, thay vì chỉ "
        "đo accuracy trên payload rõ."
    )
    doc.add_heading("2.4. Các nghiên cứu liên quan và khoảng trống", level=2)
    add_table(doc, ["Công trình", "Ý tưởng liên quan", "Ảnh hưởng đến dự án"], [
        ("DeepWAF (2019) [8]", "So sánh CNN, LSTM, CNN–LSTM và LSTM–CNN cho web attack", "Dùng CNN/LSTM đơn làm ablation và kiểm tra thứ tự ghép"),
        ("Tadhani et al. (2024) [9]", "Hybrid CNN–LSTM trên payload, CSIC 2010 và testbed SQLi/XSS", "Cơ sở chọn ba góc nhìn dữ liệu và mô hình lai; dự án tái đánh giá bằng pipeline riêng"),
        ("eXpose; URLNet [10], [12]", "Học trực tiếp trên chuỗi ký tự/URL thô", "Ủng hộ embedding ký tự và giảm feature engineering thủ công"),
        ("Rong et al. (2019) [11]", "Character-level CNN cho malicious HTTP request", "Ủng hộ giữ quan hệ giữa ký tự và motif cục bộ của query"),
        ("Tekerek (2021) [13]", "CNN phát hiện web attack trên CSIC2010v2", "Xác nhận CSIC là benchmark phù hợp cho detector HTTP"),
        ("WAF-A-MoLE; AdvSQLi [14], [15]", "Biến đổi cú pháp giữ ngữ nghĩa để né WAF", "Động lực xây OBFU và split unseen seed/kỹ thuật"),
        ("Kapoor & Narayanan (2023) [16]", "Data leakage làm kết luận ML quá lạc quan", "Tokenizer chỉ fit train, deduplicate và tách seed/kỹ thuật"),
    ])
    doc.add_paragraph(
        "Công trình gần nhất với hướng triển khai của nhóm là Tadhani et al. [9]: "
        "nghiên cứu này cũng dùng CNN–LSTM, HTTP CSIC 2010, một nguồn payload SQLi/XSS "
        "và testbed tự xây. Dự án hiện tại không sao chép số liệu của paper; toàn bộ "
        "metric trong Chương 8 được tính lại từ checkpoint của repo, cùng preprocessing "
        "và threshold 0,5. Phần mở rộng chính của nhóm là bảo toàn dấu vết encoding, "
        "hard negative, test unseen_seed/unseen_technique và kiểm tra topology tuần tự–song song."
    )
    doc.add_heading("2.5. Vì sao dùng character-level", level=2)
    doc.add_paragraph(
        "Word-level tokenizer dễ vỡ khi attacker chèn comment, ký tự phân tách hoặc mã hóa từng ký tự. Character-level giữ được cấu trúc `%27`, `&#x3c;`, dấu nháy, ngoặc, toán tử và thẻ HTML; vocabulary nhỏ, không phụ thuộc từ điển SQL/JavaScript cố định. eXpose học trực tiếp trên các chuỗi bảo mật thô bằng character embedding và CNN [10]; Rong et al. áp dụng character-level CNN cho malicious web request [11]; URLNet cũng cho thấy biểu diễn ký tự giúp xử lý token hiếm và chuỗi URL chưa thấy [12]. Các kết quả đó tạo cơ sở học thuật trực tiếp cho lựa chọn character-level của dự án. Đổi lại, sequence dài hơn và chi phí LSTM cao hơn."
    )
    doc.add_heading("2.6. Vai trò của CNN và LSTM", level=2)
    doc.add_paragraph(
        "CNN một chiều quét các cửa sổ cục bộ và phù hợp để học motif giống character n-gram. Các nghiên cứu CNN cho phân loại chuỗi cho thấy convolution kết hợp pooling có thể trích đặc trưng phân biệt hiệu quả [5], [18]. LSTM được thiết kế để duy trì và điều tiết thông tin dài hạn trong chuỗi [4], phù hợp để học thứ tự giữa breakout, keyword, operator, comment và context HTTP. DeepWAF [8] và Tadhani et al. [9] là các dẫn chứng trực tiếp trong miền web attack cho việc kết hợp năng lực trích đặc trưng của CNN với mô hình hóa phụ thuộc chuỗi của LSTM."
    )
    add_table(doc, ["Kiến trúc", "Điểm mạnh", "Hạn chế trong bài toán"], [
        ("CNN-only", "Nhanh; học motif cục bộ; song song hóa tốt", "Khó biểu diễn quan hệ xa và thứ tự toàn request"),
        ("LSTM-only", "Học thứ tự và ngữ cảnh chuỗi", "Chậm; dễ bị chuỗi dài/padding làm loãng dấu hiệu ngắn"),
        ("CNN–LSTM", "CNN lọc motif, LSTM kết hợp motif theo ngữ cảnh", "Nhiều tham số hơn; lợi thế phải được chứng minh bằng ablation"),
    ], proposed_col=0)

    doc.add_heading("3. QUÁ TRÌNH NGHIÊN CỨU VÀ CÁC QUYẾT ĐỊNH", level=1)
    add_figure(doc, "research_timeline.png", "Hình 3.1. Tiến trình phát triển dự án theo lịch sử Git", width=6.7)
    doc.add_heading("3.1. Giai đoạn khởi tạo", level=2)
    doc.add_paragraph(
        "Phiên bản đầu tập trung vào mô hình CNN–LSTM và WebApp. Sau đó nhóm bổ sung báo cáo, thử siêu tham số và xây dựng LSTM-only để trả lời câu hỏi liệu thành phần tuần tự có đủ cho bài toán hay không. CNN-only và mô hình song song được bổ sung như các thí nghiệm ablation/đối chứng."
    )
    doc.add_heading("3.2. Phát hiện vấn đề dữ liệu", level=2)
    doc.add_paragraph(
        "Dataset obfuscation thế hệ đầu cho điểm gần tuyệt đối ngay cả với Logistic Regression + TF–IDF character n-gram. Phân tích cho thấy shortcut theo nguồn sinh, ký tự đặc biệt chỉ xuất hiện ở Attack, số seed ít và seed test trùng train. Kết quả cao khi đó chủ yếu phản ánh khả năng nhận diện dấu vết generator, chưa chứng minh khả năng phát hiện obfuscation mới."
    )
    doc.add_heading("3.3. Thiết kế lại dataset và pipeline", level=2)
    doc.add_paragraph(
        "Nhóm xây dựng kho seed SQLi/XSS/benign lớn hơn, đặt Attack vào khung request benign, thêm hard negative có từ khóa/ký tự nhạy cảm, áp dụng obfuscation cho một phần benign và tạo split theo hai trục seed/kỹ thuật. Đồng thời preprocessing được gom về một module dùng chung và thêm `[USER_AGENT]` để train-serving nhất quán."
    )
    doc.add_heading("3.4. Đồng bộ thí nghiệm", level=2)
    doc.add_paragraph(
        "Trong giai đoạn hoàn thiện, nhóm kiểm tra hash `(payload, label)` của các split, đồng bộ max_len/epoch cho OBFU, rerun CNN–LSTM Kaggle/CSIC bằng preprocessing hiện tại và thay việc đọc metric metadata bằng re-score checkpoint tại threshold 0,5. Quy trình này tách rõ hai vấn đề: trọng số mô hình và quy tắc quyết định. Thay threshold không cần retrain, nhưng phải chạy inference lại."
    )

    doc.add_heading("4. DỮ LIỆU VÀ XÂY DỰNG TESTBED", level=1)
    add_figure(doc, "dataset_distribution.png", "Hình 4.1. Phân bố nhãn sau làm sạch", width=6.4)
    datasets = data["preprocessing"]["datasets"]
    dataset_rows = []
    for key, label in [("kaggle", "SQLi–XSS Payload"), ("csic", "HTTP CSIC 2010"), ("obfu_http", "OBFU_HTTP v2")]:
        clean = datasets[key]["clean"]
        dataset_rows.append((label, num(clean["rows"]), num(clean["label_counts"]["0"]), num(clean["label_counts"]["1"]), f"{clean['length']['p95']:.0f}", num(clean["length"]["max"])))
    add_caption(doc, "Bảng 4.1. Quy mô ba nguồn dữ liệu")
    add_table(doc, ["Nguồn", "Tổng", "Normal", "Attack", "P95 độ dài", "Max"], dataset_rows)
    doc.add_paragraph(
        "Ba nguồn không được chọn chỉ vì sẵn có mà để trả lời ba câu hỏi khác nhau: "
        "Kaggle kiểm tra nhận diện trực tiếp trên payload; CSIC kiểm tra request HTTP có "
        "nhiều trường; OBFU kiểm tra biến đổi che giấu và khả năng khái quát. Cách dùng "
        "nhiều nguồn tương đồng với thiết kế của Tadhani et al. [9], trong khi nghiên cứu "
        "về robustness dưới dataset shift nhấn mạnh rằng đánh giá trên nhiều phân phối độc "
        "lập cung cấp bằng chứng tốt hơn một test set duy nhất [17]."
    )
    doc.add_heading("4.1. Kaggle SQLi–XSS Payload", level=2)
    doc.add_paragraph(
        "Nguồn Kaggle cung cấp payload/chuỗi và nhãn SQLi/XSS. Sau làm sạch còn 151.663 mẫu. Nguồn payload-only giúp tách khả năng nhận diện cú pháp tấn công khỏi metadata của request, đồng thời tạo điểm đối chiếu với nghiên cứu dùng SQL/XSS payload dataset [9]. Do đây không phải raw request hoàn chỉnh, pipeline bọc payload vào HTTP template xác định, độc lập với nhãn, để có cùng không gian đầu vào với CSIC và OBFU."
    )
    doc.add_heading("4.2. HTTP CSIC 2010", level=2)
    doc.add_paragraph(
        "CSIC cung cấp request có method, URL, body, cookie và content type. Đây là benchmark được xây cho nghiên cứu web attack và đã được sử dụng trong các mô hình CNN cũng như hybrid deep learning [9], [13]. Pipeline tách path/query và serialize các trường thay vì chỉ lấy value. Sau làm sạch còn 61.065 mẫu, trong đó 36.000 Normal và 25.065 Attack; class weight được tính để cân bằng đóng góp loss."
    )
    doc.add_heading("4.3. OBFU_HTTP v2", level=2)
    stats = data["obfu_stats"]
    doc.add_paragraph(
        f"Testbed chuyên biệt có {num(stats['rows'])} request, gồm {num(stats['normal'])} Normal và {num(stats['attack'])} Attack; Attack chia đều SQLi/XSS ({num(stats['sqli'])}/{num(stats['xss'])}). Dataset hiện có {num(stats['seed_count'])} seed Attack, {num(stats['combo_count'])} tổ hợp được gán và {num(stats['technique_string_count'])} chuỗi mô tả kỹ thuật khác nhau."
    )
    add_caption(doc, "Bảng 4.2. Thiết kế split OBFU_HTTP v2")
    obfu_split_rows = []
    for split_name, split in datasets["obfu_http"]["splits"].items():
        obfu_split_rows.append((split_name, num(split["rows"]), num(split["label_counts"]["0"]), num(split["label_counts"]["1"]), f"{split['length']['p95']:.1f}".replace(".", ",")))
    add_table(doc, ["Split", "Tổng", "Normal", "Attack", "P95 độ dài"], obfu_split_rows)
    doc.add_paragraph(
        "Ba split đặc biệt kiểm tra khái quát: test_unseen_technique giữ payload seed quen nhưng kỹ thuật mã hóa lạ; test_unseen_seed giữ kỹ thuật quen nhưng seed payload lạ; test_unseen_both giữ lại cả hai. Đây là thiết kế mạnh hơn random row split vì đo đúng khả năng đối phó biến thể chưa xuất hiện trong train."
    )
    add_caption(doc, "Bảng 4.3. Vị trí mang payload trong OBFU_HTTP v2")
    context_rows = [(str(key), num(value)) for key, value in stats["context"].items()]
    add_table(doc, ["Context", "Số mẫu"], context_rows)
    doc.add_heading("4.4. Quality gate và hard negative", level=2)
    doc.add_paragraph(
        "Quality gate kiểm tra cân bằng lớp, rows/seed, seed overlap, technique holdout, shortcut header, ký tự đặc biệt trong benign, tính hợp lệ payload, baseline tuyến tính và trùng lặp với Kaggle. Dataset v2 chủ động chứa benign có apostrophe, SQL word, HTML/JS snippet, JSON quote, angle text và double encoding để mô hình không thể dùng một ký tự đơn làm luật phân loại. Việc tách các mẫu liên quan và ngăn overlap được thúc đẩy bởi bằng chứng rằng leakage có thể làm kết quả ML lạc quan nghiêm trọng [16]. Một tiêu chí baseline in-domain vẫn khó đạt vì SQLi/XSS ở mức payload vốn có tín hiệu bề mặt mạnh; hạn chế này được ghi nhận thay vì che giấu."
    )

    doc.add_heading("5. PIPELINE TIỀN XỬ LÝ", level=1)
    doc.add_heading("5.1. HTTP envelope thống nhất", level=2)
    add_code(doc, "[METHOD] ... [PATH] ... [QUERY] ... [BODY] ... [COOKIE] ... [CONTENT_TYPE] ... [USER_AGENT] ...")
    doc.add_paragraph(
        "Một biểu diễn chung giúp cùng kiến trúc xử lý payload-only và request đầy đủ. Marker trường cung cấp ranh giới ngữ nghĩa cho mô hình nhưng không tiết lộ nhãn. WebApp gọi đúng `preprocess_inference_input()` để biến payload, URL hoặc raw HTTP request thành envelope giống lúc train."
    )
    doc.add_heading("5.2. Chính sách bảo toàn obfuscation", level=2)
    doc.add_paragraph(
        "Một số nghiên cứu chuẩn hóa bằng cách decode payload trước khi phân loại [9]. "
        "Dự án chọn mục tiêu khác: detector phải quan sát chính biểu diễn mà attacker gửi "
        "và học cả dấu vết single/double encoding. Quyết định này phù hợp với threat model "
        "của WAF-A-MoLE và AdvSQLi, nơi biến đổi cú pháp là công cụ né bộ lọc [14], [15]. "
        "Do đó preprocessing chỉ chuẩn hóa định dạng thừa, không xóa tín hiệu obfuscation."
    )
    add_table(doc, ["Thao tác", "Thiết lập", "Lý do"], [
        ("URL decode", "Không", "Giữ `%27`, `%3C`, double encoding"),
        ("HTML unescape", "Không", "Giữ `&#x...;` và entity"),
        ("Lowercase", "Không", "Giữ mixed-case evasion"),
        ("Whitespace", "Chỉ chuẩn hóa dư thừa", "Giữ cấu trúc, giảm nhiễu định dạng"),
        ("Deduplicate", "Theo payload + label", "Giảm lặp; xử lý conflict nhãn"),
    ])
    doc.add_heading("5.3. Tokenizer và vector hóa", level=2)
    doc.add_paragraph(
        "Tokenizer Keras dùng `char_level=True`, `lower=False`, `filters=''`, có `<OOV>` và chỉ fit trên train split của từng nguồn. Chuỗi token được post-padding/post-truncation. Kaggle/CSIC dùng max_len 1024; OBFU dùng 768 vì P99 khoảng 671. Việc chỉ fit trên train ngăn validation/test ảnh hưởng vocabulary."
    )
    doc.add_heading("5.4. Chia dữ liệu và tái lập", level=2)
    doc.add_paragraph(
        "Kaggle và CSIC hiện dùng stratified random row split với seed 42; OBFU dùng split được đóng gói sẵn. Tokenizer chỉ fit train và checkpoint chỉ được chọn bằng validation loss. Kapoor và Narayanan [16] chỉ ra rằng phụ thuộc giữa train/test hoặc preprocessing dùng thông tin ngoài train có thể làm kết luận quá lạc quan. Vì vậy OBFU tách theo seed/kỹ thuật trước, còn random row split của Kaggle/CSIC được ghi rõ là hạn chế vì có thể để request family tương tự xuất hiện ở nhiều split."
    )

    doc.add_heading("6. LỰA CHỌN VÀ THIẾT KẾ MÔ HÌNH", level=1)
    doc.add_heading("6.1. Baseline CNN-only", level=2)
    doc.add_paragraph(
        "CNN-only gồm Embedding → Conv1D(k=3) → Pool → Conv1D(k=5) → Pool → GlobalMaxPooling → Dense → Sigmoid. Baseline này trả lời câu hỏi motif cục bộ đã đủ hay chưa và cung cấp mốc về tốc độ/độ phức tạp."
    )
    doc.add_heading("6.2. Baseline LSTM-only", level=2)
    doc.add_paragraph(
        "LSTM-only gồm Embedding → LSTM(128, return_sequences=True) → GlobalMaxPooling → Dense → Sigmoid. Baseline đo giá trị của ngữ cảnh tuần tự khi không có convolution lọc trước."
    )
    doc.add_heading("6.3. Vì sao chọn CNN–LSTM thay vì một nhánh", level=2)
    doc.add_paragraph(
        "Lựa chọn này kế thừa một hướng nghiên cứu đã có trong web security: DeepWAF "
        "thử CNN, LSTM và các cách ghép hai mạng [8]; Tadhani et al. dùng hybrid CNN–LSTM "
        "để kết hợp trích đặc trưng với phụ thuộc tuần tự trên SQLi/XSS [9]. Tuy nhiên "
        "dự án vẫn huấn luyện CNN-only và LSTM-only trên chính dữ liệu hiện tại vì kết quả "
        "từ paper không thể thay thế ablation nội bộ."
    )
    add_numbered(doc, [
        "Obfuscation tạo cả motif ngắn lẫn quan hệ dài: `%27`, `<scr`, comment là cục bộ; breakout–keyword–comment hoặc marker HTTP là quan hệ theo thứ tự.",
        "CNN làm bộ trích đặc trưng cục bộ và giảm chiều dài trước LSTM, nhờ đó LSTM không phải xử lý toàn bộ 768/1024 timestep thô.",
        "LSTM đặt activation CNN vào ngữ cảnh, hạn chế việc coi một ký tự nhạy cảm trong hard-negative benign là đủ để kết luận Attack.",
        "Kiến trúc lai vẫn nhỏ (khoảng 252–260 nghìn tham số), phù hợp demo local và dễ đóng gói thành một checkpoint.",
        "CNN-only/LSTM-only được giữ lại như ablation. Việc chọn mô hình lai dựa trên giả thuyết, paper liên quan và kết quả train lại của chính repo; không mặc định rằng mô hình phức tạp luôn có accuracy cao nhất.",
    ])
    add_figure(doc, "cnn_lstm_architecture.png", "Hình 6.1. Kiến trúc CNN–LSTM tuần tự", width=7.0)
    doc.add_heading("6.4. Giải thích từng quyết định kiến trúc", level=2)
    add_table(doc, ["Khối", "Cấu hình", "Lý do"], [
        ("Embedding", "64 chiều", "Học biểu diễn ký tự đặc thù domain; vocabulary nhỏ"),
        ("Conv1D 1", "128 filter, k=3", "Motif ngắn, encoding và toán tử"),
        ("MaxPool 1", "pool=4", "Giảm chiều dài và giữ activation mạnh"),
        ("Conv1D 2", "128 filter, k=5", "Mẫu dài hơn trên feature map"),
        ("MaxPool 2", "pool=4", "Tổng mức nén 16× trước LSTM"),
        ("LSTM", "128, return_sequences", "Giữ trạng thái ngữ cảnh tại mọi vị trí"),
        ("GlobalMaxPool", "theo thời gian", "Không phụ thuộc trạng thái cuối sau padding; lấy bằng chứng mạnh nhất"),
        ("Dense", "64, ReLU", "Kết hợp feature thành decision boundary"),
        ("Dropout", "0,3", "Regularization"),
        ("Output", "1, Sigmoid", "Xác suất Attack cho phân loại nhị phân"),
    ])
    doc.add_heading("6.5. Vì sao chọn tuần tự thay vì song song", level=2)
    doc.add_paragraph(
        "Hai topology trả lời hai giả thuyết khác nhau. Ở topology song song, CNN và "
        "LSTM cùng đọc embedding ký tự thô; hai vector đặc trưng chỉ gặp nhau ở phép "
        "concatenate. Ở topology tuần tự, Conv1D và MaxPool biến chuỗi ký tự thành chuỗi "
        "motif bậc cao, sau đó LSTM học thứ tự giữa các motif đó. Đây là đúng với giả "
        "thuyết của đề tài: dấu hiệu cục bộ như `%27`, `<scr`, comment cần được nhận diện "
        "trước, rồi mới đặt vào quan hệ breakout–keyword–operator–comment. Các công trình "
        "DeepWAF và hybrid SQLi/XSS cũng cung cấp tiền lệ cho cách ghép CNN/LSTM theo "
        "chuỗi xử lý [8], [9]."
    )
    add_numbered(doc, [
        "Tạo feature hierarchy rõ ràng: LSTM nhận activation đã mang nghĩa motif thay vì hai nhánh tự học trùng lặp từ ký tự thô.",
        "Hai lớp MaxPool(4) giảm chiều thời gian khoảng 16 lần trước recurrent layer; vì vậy phần LSTM không phải duyệt đủ 768/1024 timestep. Đây là lợi thế độ phức tạp theo thiết kế, chưa được coi là benchmark latency khi chưa đo thời gian thực nghiệm.",
        "Không cần phép concatenate hai vector và không phải cân bằng mức đóng góp của hai nhánh tại tầng Dense; đường suy luận CNN→LSTM→classifier dễ mô tả, kiểm thử và đóng gói hơn.",
        "Artifact OBFU tuần tự đã có đầy đủ test, unseen_seed, unseen_technique, unseen_both và đã được nối với WebApp; mô hình song song hiện chỉ có đối chứng Kaggle/CSIC.",
        "Quyết định cuối cùng kết hợp cơ sở kiến trúc, kết quả thực nghiệm và mức hoàn thiện triển khai; không dựa vào việc chọn hàng có điểm cao nhất trên một dataset duy nhất.",
    ])
    topology = data["sequential_parallel"]
    topology_rows = []
    for _, row in topology.iterrows():
        topology_rows.append((
            row["Dataset"], row["Architecture"], num(row["Parameters"]),
            f"{row['Accuracy']:.3f}%".replace(".", ","),
            f"{row['Attack Precision']:.3f}%".replace(".", ","),
            f"{row['Attack Recall']:.3f}%".replace(".", ","),
            f"{row['Attack F1']:.3f}%".replace(".", ","),
        ))
    add_caption(doc, "Bảng 6.1. Đối chứng CNN–LSTM tuần tự và CNN ∥ LSTM song song")
    add_table(
        doc,
        ["Dataset", "Topology", "Tham số", "Accuracy", "Precision", "Recall", "F1"],
        topology_rows,
        proposed_col=1,
    )
    add_figure(
        doc,
        "sequential_vs_parallel.png",
        "Hình 6.2. So sánh topology tuần tự và song song tại threshold 0,5",
        width=6.8,
    )
    seq_kaggle = topology[(topology["Dataset"] == "KAGGLE") & topology["Architecture"].str.contains("tuần tự")].iloc[0]
    par_kaggle = topology[(topology["Dataset"] == "KAGGLE") & topology["Architecture"].str.contains("song song")].iloc[0]
    seq_csic = topology[(topology["Dataset"] == "CSIC") & topology["Architecture"].str.contains("tuần tự")].iloc[0]
    par_csic = topology[(topology["Dataset"] == "CSIC") & topology["Architecture"].str.contains("song song")].iloc[0]
    doc.add_paragraph(
        f"Sau lần train đồng bộ, tuần tự cao hơn song song "
        f"{seq_kaggle['Attack F1'] - par_kaggle['Attack F1']:.3f} điểm phần trăm F1 "
        f"trên Kaggle ({seq_kaggle['Attack F1']:.3f}% so với "
        f"{par_kaggle['Attack F1']:.3f}%), nhưng thấp hơn "
        f"{par_csic['Attack F1'] - seq_csic['Attack F1']:.3f} điểm trên CSIC "
        f"({seq_csic['Attack F1']:.3f}% so với {par_csic['Attack F1']:.3f}%). "
        "Do đó dữ liệu không chứng minh tuần tự vượt trội phổ quát. Nhóm giữ tuần tự "
        "vì nó phù hợp giả thuyết motif→ngữ cảnh, tốt hơn trên nguồn payload lớn và là "
        "topology đã hoàn tất đánh giá OBFU/triển khai."
    )
    doc.add_heading("6.6. BiLSTM và các hướng mở rộng", level=2)
    doc.add_paragraph(
        "Repo còn thí nghiệm CNN–BiLSTM để khai thác ngữ cảnh hai chiều. BiLSTM phù hợp "
        "khi toàn chuỗi đã sẵn sàng trước suy luận, nhưng tăng tham số và chi phí. Vì "
        "mục tiêu hiện tại là một detector nhị phân gọn để tích hợp WebApp, BiLSTM được "
        "giữ như hướng mở rộng thay vì thay checkpoint triển khai."
    )

    doc.add_heading("7. HUẤN LUYỆN VÀ ĐÁNH GIÁ", level=1)
    doc.add_heading("7.1. Cấu hình huấn luyện", level=2)
    kaggle_model = data["models"]["CNN–LSTM"]["kaggle"]["model"]
    csic_model = data["models"]["CNN–LSTM"]["csic"]["model"]
    obfu_model = data["models"]["CNN–LSTM"]["obfu_http"]["model"]
    add_table(doc, ["Thành phần", "Kaggle", "CSIC", "OBFU_HTTP"], [
        ("Seed", "42", "42", "42"),
        ("Max length", "1024", "1024", "768"),
        ("Embedding", "64", "64", "64"),
        ("Batch size", "128", "128", "128"),
        ("Epoch tối đa", str(kaggle_model["epochs_requested"]), str(csic_model["epochs_requested"]), str(obfu_model["epochs_requested"])),
        ("Epoch thực chạy", str(kaggle_model["epochs_ran"]), str(csic_model["epochs_ran"]), str(obfu_model["epochs_ran"])),
        ("Số tham số", num(kaggle_model["parameter_count"]), num(csic_model["parameter_count"]), num(obfu_model["parameter_count"])),
        ("Optimizer", "Adam", "Adam", "Adam"),
        ("Learning rate", "0,001", "0,001", "0,001"),
        ("Loss", "Binary cross-entropy", "Binary cross-entropy", "Binary cross-entropy"),
        ("Early stopping", "val_loss, patience=3", "val_loss, patience=3", "val_loss, patience=3"),
        ("Threshold", "0,5", "0,5", "0,5"),
    ])
    doc.add_paragraph(
        "Lần train mới nhất dừng sớm sau 12 epoch trên Kaggle và 17 epoch trên CSIC; "
        "OBFU chạy 11 epoch. Đây là số epoch thực tế được đọc trực tiếp từ metadata "
        "checkpoint, không phải giả định rằng cả ba mô hình đều chạy đến epoch tối đa."
    )
    doc.add_heading("7.2. Class weight", level=2)
    doc.add_paragraph(
        "Class weight được tính `balanced` trên train split. OBFU cân bằng nên trọng số hai lớp bằng 1. Kaggle có nhiều Attack hơn nên Normal được tăng trọng số; CSIC có nhiều Normal hơn nên Attack được tăng trọng số. Class weight tác động loss trong train, không thay thế việc đánh giá Precision/Recall theo từng lớp."
    )
    doc.add_heading("7.3. Metric", level=2)
    add_bullets(doc, [
        "Accuracy: tỷ lệ dự đoán đúng trên toàn bộ mẫu.",
        "Precision Attack = TP/(TP+FP): độ tin cậy của cảnh báo Attack.",
        "Recall Attack = TP/(TP+FN): khả năng không bỏ sót tấn công; đặc biệt quan trọng trong an ninh.",
        "F1 = 2·Precision·Recall/(Precision+Recall): cân bằng Precision và Recall.",
        "AUC-ROC/PR-AUC: đánh giá xếp hạng xác suất, ít phụ thuộc một threshold cụ thể.",
    ])
    doc.add_heading("7.4. Nguyên tắc so sánh công bằng", level=2)
    doc.add_paragraph(
        "Ba mô hình phải dùng cùng `(payload,label)` test split, cùng tokenizer policy, max_len tương ứng, cùng seed và cùng threshold. Báo cáo tạo bảng bằng cách nạp best checkpoint rồi inference lại tại 0,5; không lấy trực tiếp metric val-F1 cũ trong metadata. Số epoch thực chạy khác nhau là hợp lệ vì cùng early stopping; số tham số khác nhau là đặc điểm của kiến trúc cần so sánh."
    )

    doc.add_heading("8. KẾT QUẢ THỰC NGHIỆM", level=1)
    doc.add_heading("8.1. So sánh CNN, LSTM và CNN–LSTM", level=2)
    comparison = data["comparison"]
    comparison_rows = []
    for _, row in comparison.iterrows():
        model_name = "CNN–LSTM" if row["DL Model"] == "PROPOSED MODEL" else row["DL Model"]
        comparison_rows.append((model_name, row["Dataset"], f"{row['Accuracy']:.2f}%".replace(".", ","), f"{row['Precision']:.2f}%".replace(".", ","), f"{row['Recall']:.2f}%".replace(".", ","), f"{row['F1-Score']:.2f}%".replace(".", ",")))
    add_caption(doc, "Bảng 8.1. Kết quả re-score tại threshold 0,5")
    add_table(doc, ["Mô hình", "Dataset", "Accuracy", "Precision", "Recall", "F1"], comparison_rows, proposed_col=0)
    add_figure(doc, "fair_model_comparison.png", "Hình 8.1. So sánh ba kiến trúc tại threshold 0,5", width=7.0)
    cnn_kaggle = comparison[(comparison["DL Model"] == "CNN") & (comparison["Dataset"] == "SQLi - XSS Payload")].iloc[0]
    lstm_kaggle = comparison[(comparison["DL Model"] == "LSTM") & (comparison["Dataset"] == "SQLi - XSS Payload")].iloc[0]
    hybrid_kaggle_result = comparison[(comparison["DL Model"] == "PROPOSED MODEL") & (comparison["Dataset"] == "SQLi - XSS Payload")].iloc[0]
    cnn_csic = comparison[(comparison["DL Model"] == "CNN") & (comparison["Dataset"] == "HTTP CSIC2010")].iloc[0]
    lstm_csic = comparison[(comparison["DL Model"] == "LSTM") & (comparison["Dataset"] == "HTTP CSIC2010")].iloc[0]
    hybrid_csic_result = comparison[(comparison["DL Model"] == "PROPOSED MODEL") & (comparison["Dataset"] == "HTTP CSIC2010")].iloc[0]
    doc.add_paragraph(
        f"Sau khi train lại bằng preprocessing hiện tại và re-score ở threshold 0,5, "
        f"CNN–LSTM tuần tự đạt F1 {hybrid_kaggle_result['F1-Score']:.3f}% trên Kaggle, "
        f"cao hơn LSTM-only {lstm_kaggle['F1-Score']:.3f}% và CNN-only "
        f"{cnn_kaggle['F1-Score']:.3f}%. Trên CSIC, CNN–LSTM đạt F1 "
        f"{hybrid_csic_result['F1-Score']:.3f}%, so với LSTM "
        f"{lstm_csic['F1-Score']:.3f}% và CNN {cnn_csic['F1-Score']:.3f}%. "
        "Kết quả mới không còn hiện tượng suy giảm Kaggle của checkpoint cũ khi áp "
        "threshold 0,5, vì mô hình đã được train lại với chiến lược threshold cố định "
        "và pipeline đồng bộ."
    )
    doc.add_paragraph(
        "CNN–LSTM đứng đầu trong đối chứng ba mô hình trên Kaggle và CSIC, nhưng trên "
        "OBFU CNN-only vẫn nhỉnh hơn rất nhỏ. Ngoài ra, mô hình song song vượt tuần tự "
        "trên CSIC ở thí nghiệm topology. Vì vậy kết luận phù hợp là train lại đã xác "
        "nhận lợi ích thực nghiệm của mô hình lai trên hai nguồn chính, không phải bằng "
        "chứng rằng mọi biến thể CNN–LSTM luôn vượt mọi baseline trên mọi phân phối."
    )
    doc.add_heading("8.2. Kết quả CNN–LSTM trên OBFU_HTTP", level=2)
    obfu_eval_rows = []
    for name, result in data["proposed_obfu"]["evaluation"].items():
        label = name.split(":")[-1] if ":" in name else "test"
        m = metric_row(result)
        obfu_eval_rows.append((label, pct(m["accuracy"]), pct(m["precision"]), pct(m["recall"]), pct(m["f1"]), str(m["cm"])))
    add_caption(doc, "Bảng 8.2. CNN–LSTM trên testbed OBFU_HTTP v2")
    add_table(doc, ["Split", "Accuracy", "Precision", "Recall", "F1", "Confusion matrix"], obfu_eval_rows)
    add_figure(doc, "obfu_generalization.png", "Hình 8.2. Khả năng khái quát trên split OBFU", width=6.8)
    main_result = metric_row(data["proposed_obfu"]["evaluation"]["obfu_http"])
    unseen_result = metric_row(data["proposed_obfu"]["evaluation"]["obfu_http:test_unseen_both"])
    doc.add_paragraph(
        f"Trên test chuẩn, model chỉ có {main_result['cm'][0][1]} false positive và {main_result['cm'][1][0]} false negative. Trên unseen_both, có {unseen_result['cm'][0][1]} false positive và {unseen_result['cm'][1][0]} false negative. Mức suy giảm F1 nhỏ cho thấy model học được nhiều dấu hiệu vượt ra ngoài seed/kỹ thuật quen trong testbed hiện tại."
    )
    doc.add_heading("8.3. Diễn giải thận trọng", level=2)
    add_bullets(doc, [
        "Kết quả gần 100% không đồng nghĩa bảo vệ được mọi ứng dụng thực tế; dataset vẫn là dữ liệu tổng hợp/công khai có phạm vi hữu hạn.",
        "OBFU v2 giảm shortcut và có held-out design, nhưng baseline tuyến tính in-domain vẫn cao; nhãn Attack có tín hiệu bề mặt mạnh.",
        "Kaggle/CSIC dùng random row split nên có thể lạc quan hơn group split.",
        "Chênh lệch nhỏ giữa ba kiến trúc chưa đủ để kết luận mô hình lai tốt hơn có ý nghĩa thống kê; cần nhiều seed và kiểm định cặp.",
        "Metric phải được đọc cùng confusion matrix và Recall Attack, không chỉ Accuracy.",
    ])

    doc.add_heading("9. TRIỂN KHAI WEBAPP", level=1)
    add_figure(doc, "webapp_flow.png", "Hình 9.1. Luồng suy luận end-to-end", width=6.8)
    doc.add_heading("9.1. Kiến trúc triển khai", level=2)
    add_table(doc, ["Thành phần", "Công nghệ / vai trò"], [
        ("Frontend", "HTML, CSS, JavaScript; nhập payload và hiển thị xác suất"),
        ("Backend", "Flask; route, validation, JSON API"),
        ("Preprocessing", "preprocess_inference_input(); nhận envelope/raw HTTP/URL/payload"),
        ("Tokenizer", "tokenizer.pkl; character-level giống train"),
        ("Model", "best_hybrid_cnn_lstm.keras, load compile=False"),
        ("Metadata", "max_len và cấu hình artifact"),
    ])
    doc.add_heading("9.2. Backend Flask", level=2)
    doc.add_paragraph(
        "`webapp/app.py` trỏ tới model OBFU tại `cnn_lstm/artifacts_cnn_lstm_by_dataset/by_dataset/obfu_http`. Model và tokenizer được lazy-load trong request dự đoán đầu tiên và cache bằng `lru_cache(maxsize=1)`, tránh tải lại cho mỗi request. Input tối đa 100.000 ký tự; threshold phải thuộc [0,1]. Flask phù hợp prototype vì cung cấp routing, template và JSON API nhẹ [7]."
    )
    add_caption(doc, "Bảng 9.1. Endpoint chính")
    add_table(doc, ["Method", "Endpoint", "Chức năng"], [
        ("GET", "/", "Giao diện WebApp"),
        ("GET", "/api/health", "Kiểm tra model, tokenizer và runtime"),
        ("POST", "/api/predict", "Chuẩn hóa input, vector hóa và dự đoán"),
    ])
    doc.add_heading("9.3. Train-serving consistency", level=2)
    doc.add_paragraph(
        "Đây là lỗi quan trọng từng được phát hiện trong quá trình nghiên cứu: model train trên HTTP envelope nhưng webapp ban đầu gửi chuỗi thô. Phiên bản hiện tại đã sửa bằng cách dùng chung `preprocess_inference_input()`. Hàm tự nhận dạng bốn loại input: envelope đã chuẩn hóa, raw HTTP request, URL/path và payload thuần. Kết quả API trả cả `normalized_payload`, `input_kind`, token count, max_len và cờ truncated để người dùng kiểm chứng đầu vào thực tế của model."
    )
    doc.add_heading("9.4. Luồng giao diện", level=2)
    add_numbered(doc, [
        "Người dùng nhập chuỗi và threshold (mặc định 0,5).",
        "JavaScript gửi JSON đến POST /api/predict bằng fetch.",
        "Backend validate, chuẩn hóa và vector hóa.",
        "Model trả xác suất Attack; backend áp threshold.",
        "Frontend hiển thị nhãn, Attack/Normal probability, input kind, độ dài và chuỗi đã chuẩn hóa.",
    ])
    add_code(doc, '{\n  "payload": "/search?q=\' OR 1=1 --",\n  "threshold": 0.5\n}')
    doc.add_heading("9.5. Vận hành", level=2)
    add_code(doc, "cd webapp\npython app.py\n# Truy cập http://127.0.0.1:8000")
    doc.add_paragraph(
        "Server tích hợp của Flask phù hợp demo/local test, không phải production server. Nếu triển khai thật cần WSGI server, TLS/reverse proxy, authentication, rate limiting, logging, monitoring, model versioning và chính sách bảo vệ dữ liệu request."
    )

    doc.add_heading("10. ĐÁNH GIÁ HỆ THỐNG VÀ HẠN CHẾ", level=1)
    doc.add_heading("10.1. Điểm đạt được", level=2)
    add_bullets(doc, [
        "Pipeline từ dataset → preprocessing → train → artifact → API → giao diện đã khép kín.",
        "Có ba baseline/ablation và quy tắc threshold chung cho bảng so sánh.",
        "OBFU v2 có hard negative và ba split khái quát, tốt hơn random test đơn thuần.",
        "WebApp xử lý nhiều dạng input và đảm bảo train-serving consistency.",
        "Artifact, tokenizer, metadata, history và processed split được lưu tách theo dataset.",
    ])
    doc.add_heading("10.2. Hạn chế", level=2)
    add_table(doc, ["Hạn chế", "Ảnh hưởng", "Hướng xử lý"], [
        ("Random row split Kaggle/CSIC", "Có thể có family gần nhau ở train/test", "Bổ sung family-group split"),
        ("Dataset tổng hợp", "Chưa bao phủ traffic và attacker thật", "Thu thập log đã ẩn danh, red-team test"),
        ("Phân loại nhị phân", "Không biết SQLi hay XSS và subtype", "Multi-class/multi-label"),
        ("Chỉ ba seed train chính", "Chưa đo phương sai kiến trúc đầy đủ", "Nhiều seed, bootstrap, paired test"),
        ("Input bị truncate", "Payload cuối chuỗi có thể bị mất", "Sliding window/segment pooling"),
        ("WebApp prototype", "Thiếu auth/rate-limit/production hardening", "Triển khai WSGI + gateway + monitoring"),
        ("Model có thể bị adversarial evasion", "False negative dưới biến đổi mới", "Adversarial training và drift monitoring"),
    ])
    doc.add_heading("10.3. Nguyên tắc sử dụng an toàn", level=2)
    doc.add_paragraph(
        "Detector không được dùng như bằng chứng duy nhất để chặn request hoặc kết luận người dùng tấn công. Trong production nên kết hợp rule/WAF, xác thực ngữ cảnh, secure coding, tham số hóa truy vấn, output encoding, CSP, logging và cơ chế human review. Threshold cần được hiệu chỉnh theo chi phí false positive/false negative của hệ thống đích."
    )

    doc.add_heading("11. KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN", level=1)
    doc.add_heading("11.1. Kết luận", level=2)
    doc.add_paragraph(
        "Dự án đã hiện thực hóa đề tài “Ứng dụng học máy để phát hiện payload XSS, SQLi bị obfuscation” thành một hệ thống học sâu character-level hoàn chỉnh. Quá trình nghiên cứu cho thấy chất lượng dataset và hợp đồng so sánh quan trọng không kém kiến trúc. Việc xây dựng baseline CNN/LSTM, phát hiện shortcut ở dataset v1, thiết kế OBFU v2, đồng bộ preprocessing/threshold và sửa train-serving mismatch là các kết quả phương pháp luận quan trọng."
    )
    doc.add_paragraph(
        "CNN–LSTM tuần tự được chọn vì CNN học motif cục bộ và nén chuỗi trước khi LSTM học quan hệ theo thứ tự; sequence global max pooling lấy activation ngữ cảnh mạnh ở mọi vị trí. Sau lần train đồng bộ, mô hình đạt F1 99,880% trên Kaggle, 99,379% trên CSIC và 99,963% trên OBFU test. So với topology song song, tuần tự tốt hơn trên Kaggle nhưng thấp hơn nhẹ trên CSIC. Kết luận đúng là topology tuần tự phù hợp giả thuyết motif→ngữ cảnh, có kết quả cạnh tranh và đã hoàn thiện đánh giá/triển khai; chưa có bằng chứng thống kê rằng nó luôn tốt hơn song song hoặc CNN/LSTM đơn trên mọi dữ liệu thực."
    )
    doc.add_heading("11.2. Hướng phát triển", level=2)
    add_numbered(doc, [
        "Rerun nhiều seed và báo cáo trung bình ± độ lệch chuẩn; dùng paired bootstrap/McNemar.",
        "Bổ sung family-group split cho Kaggle/CSIC và cross-dataset domain shift.",
        "Thử sliding-window, attention hoặc pooling theo field để tránh truncate.",
        "Phân loại đa lớp SQLi/XSS và subtype/kỹ thuật obfuscation.",
        "Adversarial training từ các biến đổi mới; theo dõi drift và hard-negative mining.",
        "Giải thích dự đoán bằng occlusion/saliency theo ký tự và marker HTTP.",
        "Đóng gói production với Docker, WSGI, reverse proxy, TLS, rate limit và audit log.",
    ])

    doc.add_heading("TÀI LIỆU THAM KHẢO VÀ LIÊN KẾT PAPER", level=1)
    doc.add_paragraph(
        "Các liên kết màu xanh dưới đây là hyperlink trực tiếp tới DOI, trang xuất bản "
        "hoặc bản thảo arXiv/PMLR của công trình. Nguồn [8]–[18] là nhóm paper được "
        "dùng trực tiếp để lập luận cho mô hình, dữ liệu, preprocessing, obfuscation và "
        "thiết kế đánh giá của dự án."
    )
    references = [
        ("[1] OWASP Foundation, “SQL Injection” (truy cập 05/08/2026).", "https://owasp.org/www-community/attacks/SQL_Injection"),
        ("[2] OWASP Foundation, “Cross Site Scripting (XSS)” (truy cập 05/08/2026).", "https://owasp.org/www-community/attacks/xss/"),
        ("[3] OWASP Web Security Testing Guide, “Encoded Injection” (truy cập 05/08/2026).", "https://owasp.org/www-project-web-security-testing-guide/stable/6-Appendix/D-Encoded_Injection"),
        ("[4] S. Hochreiter và J. Schmidhuber, “Long Short-Term Memory”, Neural Computation, 9(8), 1735–1780, 1997.", "https://doi.org/10.1162/neco.1997.9.8.1735"),
        ("[5] Y. Kim, “Convolutional Neural Networks for Sentence Classification”, EMNLP 2014, arXiv:1408.5882.", "https://arxiv.org/abs/1408.5882"),
        ("[6] Keras, “GlobalMaxPooling1D” (truy cập 05/08/2026).", "https://keras.io/api/layers/pooling_layers/global_max_pooling1d/"),
        ("[7] Pallets Projects, “Flask Documentation – Quickstart” (truy cập 05/08/2026).", "https://flask.palletsprojects.com/en/stable/quickstart/"),
        ("[8] X. Kuang, M. Zhang, H. Li, G. Zhao, H. Cao, Z. Wu và X. Wang, “DeepWAF: Detecting Web Attacks Based on CNN and LSTM Models”, CSS 2019, tr. 121–136.", "https://doi.org/10.1007/978-3-030-37352-8_11"),
        ("[9] J. R. Tadhani, V. Vekariya, V. Sorathiya, S. Alshathri và W. El-Shafai, “Securing web applications against XSS and SQLi attacks using a novel deep learning approach”, Scientific Reports, 14, 1803, 2024.", "https://doi.org/10.1038/s41598-023-48845-4"),
        ("[10] J. Saxe và K. Berlin, “eXpose: A Character-Level Convolutional Neural Network with Embeddings for Detecting Malicious URLs, File Paths and Registry Keys”, arXiv:1702.08568, 2017.", "https://arxiv.org/abs/1702.08568"),
        ("[11] W. Rong, B. Zhang và X. Lv, “Malicious Web Request Detection Using Character-Level CNN”, ML4CS 2019, tr. 6–16.", "https://arxiv.org/abs/1811.08641"),
        ("[12] H. Le, Q. Pham, D. Sahoo và S. C. H. Hoi, “URLNet: Learning a URL Representation with Deep Learning for Malicious URL Detection”, KDD 2018, arXiv:1802.03162.", "https://arxiv.org/abs/1802.03162"),
        ("[13] A. Tekerek, “A novel architecture for web-based attack detection using convolutional neural network”, Computers & Security, 100, 102096, 2021.", "https://doi.org/10.1016/j.cose.2020.102096"),
        ("[14] L. Demetrio, A. Valenza, G. Costa và G. Lagorio, “WAF-A-MoLE: Evading Web Application Firewalls through Adversarial Machine Learning”, SAC 2020, arXiv:2001.01952.", "https://arxiv.org/abs/2001.01952"),
        ("[15] Z. Qu, X. Ling, T. Wang, X. Chen, S. Ji và C. Wu, “AdvSQLi: Generating Adversarial SQL Injections against Real-world WAF-as-a-service”, arXiv:2401.02615, 2024.", "https://arxiv.org/abs/2401.02615"),
        ("[16] S. Kapoor và A. Narayanan, “Leakage and the reproducibility crisis in machine-learning-based science”, Patterns, 4(9), 100804, 2023.", "https://doi.org/10.1016/j.patter.2023.100804"),
        ("[17] A. Subbaswamy, R. Adams và S. Saria, “Evaluating Model Robustness and Stability to Dataset Shift”, AISTATS 2021, PMLR 130:2611–2619.", "https://proceedings.mlr.press/v130/subbaswamy21a.html"),
        ("[18] X. Zhang, J. Zhao và Y. LeCun, “Character-level Convolutional Networks for Text Classification”, NeurIPS 2015, tr. 649–657.", "https://arxiv.org/abs/1509.01626"),
        ("[19] Mã nguồn, notebook, checkpoint và metadata nội bộ của dự án obfuscated-web-attack-detection, workspace ngày 05/08/2026.", None),
    ]
    for reference, url in references:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.first_line_indent = Cm(-0.7)
        paragraph.paragraph_format.left_indent = Cm(0.7)
        paragraph.add_run(reference)
        if url:
            paragraph.add_run(" ")
            add_hyperlink(paragraph, url, "[Mở liên kết]")

    doc.add_heading("PHỤ LỤC A. CẤU TRÚC THƯ MỤC", level=1)
    add_code(doc, """obfuscated-web-attack-detection/
├── preprocessing/preprocess_data.py
├── dataset_builder/              # generator, seed pools, quality gate
├── cnn_only/                     # baseline CNN
├── lstm_only/                    # baseline LSTM
├── cnn_lstm/CNN_LSTM.py          # mô hình đề xuất
├── cnn_lstm/artifacts_cnn_lstm_by_dataset/ # artifact của ba dataset
├── analysis/                     # phân tích lỗi/evasion
├── experiments/                  # BiLSTM, sequence pool
├── webapp/app.py                 # Flask backend
├── webapp/templates/index.html   # giao diện
└── webapp/static/                # JavaScript/CSS""")

    doc.add_heading("PHỤ LỤC B. LỆNH TÁI LẬP", level=1)
    add_code(doc, """# Sinh lại dataset OBFU v2
python dataset_builder/generate_obfu_dataset_v2.py --scale full

# Train CNN–LSTM OBFU
python cnn_lstm/CNN_LSTM.py --datasets obfu_http --train-sources obfu_http \
  --epochs 30 --threshold-strategy fixed

# Phân tích lỗi
python analysis/analyze_cnn_lstm.py

# Chạy WebApp
cd webapp
python app.py""")

    doc.add_heading("PHỤ LỤC C. ARTIFACT TRIỂN KHAI", level=1)
    add_table(doc, ["Artifact", "Vai trò"], [
        ("best_hybrid_cnn_lstm.keras", "Kiến trúc + trọng số checkpoint tốt nhất"),
        ("tokenizer.pkl", "Ánh xạ ký tự → token ID"),
        ("metadata_and_results.json", "max_len, cấu hình, metric và history"),
        ("training_history.csv", "loss/metric theo epoch"),
        ("processed_data_by_dataset/", "Split đã preprocess để audit/tái lập"),
    ])

    doc.add_heading("PHỤ LỤC D. LIÊN KẾT NHANH ĐẾN CÁC PAPER", level=1)
    doc.add_paragraph(
        "Danh sách này lặp lại riêng các công trình khoa học được dùng trực tiếp trong "
        "lập luận, để người đọc có thể mở paper ngay từ cuối báo cáo."
    )
    for reference, url in references[7:18]:
        paragraph = doc.add_paragraph(style="List Bullet")
        title = reference.split("“", 1)[-1].split("”", 1)[0]
        paragraph.add_run(title + " – ")
        add_hyperlink(paragraph, url, url)

    core_props = doc.core_properties
    core_props.title = "Ứng dụng học máy để phát hiện payload XSS và SQLi bị obfuscation"
    core_props.subject = "Báo cáo quá trình nghiên cứu, thực nghiệm CNN–LSTM và triển khai WebApp"
    core_props.author = "Nhóm nghiên cứu"
    core_props.keywords = "SQL Injection, XSS, obfuscation, CNN, LSTM, Flask"
    core_props.comments = "Sinh tự động từ mã nguồn và artifact ngày 05/08/2026."
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_DOCX)
    print(f"Saved report: {OUTPUT_DOCX}")


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--assets", action="store_true", help="Generate PNG charts/diagrams.")
    args = parser.parse_args()
    data = collect_data()
    if args.assets:
        generate_assets(data)
        print(f"Saved assets: {ASSET_DIR}")
    else:
        generate_docx(data)


if __name__ == "__main__":
    main()
